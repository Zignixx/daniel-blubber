import io
import unittest

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from blubber_core import CountOptions, count_document, parse_docx


def _element(tag, **attributes):
    element = OxmlElement(tag)
    for name, value in attributes.items():
        element.set(qn(f"w:{name}"), str(value))
    return element


def _add_numbering_definition(document):
    numbering = document.part.numbering_part.element

    abstract = _element("w:abstractNum", abstractNumId=99)
    abstract.append(_element("w:multiLevelType", val="multilevel"))

    for level, start, pattern in (
        (0, 3, "Kapitel %1"),
        (1, 2, "%1.%2"),
    ):
        level_element = _element("w:lvl", ilvl=level)
        level_element.append(_element("w:start", val=start))
        level_element.append(_element("w:numFmt", val="decimal"))
        level_element.append(_element("w:lvlText", val=pattern))
        abstract.append(level_element)

    instance = _element("w:num", numId=99)
    instance.append(_element("w:abstractNumId", val=99))
    numbering.append(abstract)
    numbering.append(instance)

    # Die Formatvorlagen tragen die Nummerierung direkt (so wie es echte
    # Word-Vorlagen tun) -- nicht bloss eine lose w:pStyle-Verknuepfung auf
    # Abschnitts-Ebene, die real gar nicht zum automatischen Nummerieren fuehrt.
    for level, style_name in ((0, "Heading 1"), (1, "Heading 2")):
        style_element = document.styles[style_name].element
        p_pr = style_element.find(qn("w:pPr"))
        if p_pr is None:
            p_pr = _element("w:pPr")
            style_element.append(p_pr)
        num_pr = _element("w:numPr")
        num_pr.append(_element("w:ilvl", val=level))
        num_pr.append(_element("w:numId", val=99))
        p_pr.append(num_pr)


def _add_linked_style_numbering_definition(document):
    """Bildet Word-Vorlagen nach, die Kapitelnummern per verlinkter
    Listen-Formatvorlage (numStyleLink) statt per eingebetteten <w:lvl>
    definieren -- z. B. Uni-Vorlagen wie "Hausarbeit"."""
    numbering = document.part.numbering_part.element

    real_abstract = _element("w:abstractNum", abstractNumId=10)
    for level, start, pattern in ((0, 1, "%1"), (1, 1, "%1.%2")):
        level_element = _element("w:lvl", ilvl=level)
        level_element.append(_element("w:start", val=start))
        level_element.append(_element("w:numFmt", val="decimal"))
        level_element.append(_element("w:lvlText", val=pattern))
        real_abstract.append(level_element)
    numbering.append(real_abstract)

    style_num_id = _element("w:num", numId=20)
    style_num_id.append(_element("w:abstractNumId", val=10))
    numbering.append(style_num_id)

    numbering_style = _element("w:style", type="numbering", styleId="Hausarbeit")
    style_p_pr = _element("w:pPr")
    style_num_pr = _element("w:numPr")
    style_num_pr.append(_element("w:numId", val=20))
    style_p_pr.append(style_num_pr)
    numbering_style.append(style_p_pr)
    document.styles.element.append(numbering_style)

    linked_abstract = _element("w:abstractNum", abstractNumId=11)
    num_style_link = _element("w:numStyleLink", val="Hausarbeit")
    linked_abstract.append(num_style_link)
    numbering.append(linked_abstract)

    used_num_id = _element("w:num", numId=30)
    used_num_id.append(_element("w:abstractNumId", val=11))
    numbering.append(used_num_id)

    return 30  # numId, das die Kapitel-Absaetze direkt referenzieren


def _apply_direct_numbering(paragraph, num_id, ilvl=None):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = _element("w:numPr")
    if ilvl is not None:
        num_pr.append(_element("w:ilvl", val=ilvl))
    num_pr.append(_element("w:numId", val=num_id))
    p_pr.append(num_pr)


class ChapterNumberingTests(unittest.TestCase):
    def test_uses_rendered_word_numbering_instead_of_own_counter(self):
        document = Document()
        _add_numbering_definition(document)

        document.add_heading("Einleitung", level=1)
        document.add_paragraph("Text im ersten Kapitel.")

        document.add_heading("Hintergrund", level=2)

        document.add_heading("Ergebnisse", level=1)

        data = io.BytesIO()
        document.save(data)
        data.seek(0)

        root = parse_docx(data)
        self.assertEqual(
            [(chapter.number, chapter.title) for chapter in root.children],
            [("Kapitel 3", "Einleitung"), ("Kapitel 4", "Ergebnisse")],
        )
        self.assertEqual(
            (root.children[0].children[0].number, root.children[0].children[0].title),
            ("3.2", "Hintergrund"),
        )

        counted = count_document(root, CountOptions())
        self.assertEqual(counted[0].number, "Kapitel 3")
        self.assertEqual(counted[0].children[0].number, "3.2")
        self.assertEqual(counted[1].number, "Kapitel 4")

    def test_keeps_manual_number_and_does_not_invent_one(self):
        document = Document()
        document.add_heading("7.2. Manuell nummeriert", level=1)
        document.add_heading("Ohne Nummer", level=1)

        data = io.BytesIO()
        document.save(data)
        data.seek(0)

        root = parse_docx(data)
        self.assertEqual(root.children[0].number, "7.2.")
        self.assertEqual(root.children[0].title, "Manuell nummeriert")
        self.assertEqual(root.children[1].number, "")
        self.assertEqual(root.children[1].title, "Ohne Nummer")

    def test_resolves_numbering_linked_via_numstylelink(self):
        """Regressionstest fuer eine reale Bachelorarbeit-Vorlage: die
        Kapitel-Ueberschriften referenzieren direkt einen numId, dessen
        abstractNum keine eigenen <w:lvl>-Level hat, sondern per
        numStyleLink auf eine Listen-Formatvorlage ("Hausarbeit") verweist.
        Vorher lieferte das Tool dafuer keine Nummer ("(keine)")."""
        document = Document()
        used_num_id = _add_linked_style_numbering_definition(document)

        heading = document.add_heading("Einleitung", level=1)
        _apply_direct_numbering(heading, used_num_id)

        sub_heading = document.add_heading("Relevanz und Kontext", level=2)
        _apply_direct_numbering(sub_heading, used_num_id, ilvl=1)

        document.add_heading("Forschungsstand", level=1)
        _apply_direct_numbering(document.paragraphs[-1], used_num_id)

        data = io.BytesIO()
        document.save(data)
        data.seek(0)

        root = parse_docx(data)
        self.assertEqual(root.children[0].number, "1")
        self.assertEqual(root.children[0].title, "Einleitung")
        self.assertEqual(root.children[0].children[0].number, "1.1")
        self.assertEqual(root.children[1].number, "2")
        self.assertEqual(root.children[1].title, "Forschungsstand")

    def test_headings_without_any_numpr_stay_unnumbered(self):
        """Regressionstest: eine unbenutzte Listen-Definition, die zufaellig
        per pStyle auf "Heading 1" verweist, darf Ueberschriften OHNE eigene
        Nummerierung (z. B. Verzeichnis-Seiten wie "Abbildungsverzeichnis")
        keine erfundene Nummer verpassen. Das war der urspruengliche Bug."""
        document = Document()
        numbering = document.part.numbering_part.element

        stray_abstract = _element("w:abstractNum", abstractNumId=77)
        stray_level = _element("w:lvl", ilvl=0)
        stray_level.append(_element("w:start", val=1))
        stray_level.append(_element("w:numFmt", val="decimal"))
        stray_level.append(_element("w:lvlText", val="%1"))
        stray_level.append(_element("w:pStyle", val="Heading1"))
        stray_abstract.append(stray_level)
        numbering.append(stray_abstract)

        stray_num = _element("w:num", numId=77)
        stray_num.append(_element("w:abstractNumId", val=77))
        numbering.append(stray_num)

        document.add_heading("Abbildungsverzeichnis", level=1)
        document.add_heading("Tabellenverzeichnis", level=1)

        data = io.BytesIO()
        document.save(data)
        data.seek(0)

        root = parse_docx(data)
        self.assertEqual(root.children[0].number, "")
        self.assertEqual(root.children[1].number, "")


if __name__ == "__main__":
    unittest.main()
